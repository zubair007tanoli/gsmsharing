#!/usr/bin/env python3
"""
PDF to Word conversion using pdfplumber for accurate table extraction.
This script provides better table detection and extraction compared to pdf2docx.
Requires: pip install pdfplumber python-docx
"""
import sys
import os


def main(input_pdf, output_docx):
    """Convert PDF to Word document using pdfplumber for accurate extraction."""
    try:
        import pdfplumber
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as e:
        print(f"ERROR: Install required packages: pip install pdfplumber python-docx. {e}", file=sys.stderr)
        sys.exit(1)

    try:
        if not os.path.exists(input_pdf):
            print(f"ERROR: Input file not found: {input_pdf}", file=sys.stderr)
            sys.exit(1)

        # Ensure output directory exists
        out_dir = os.path.dirname(output_docx)
        if out_dir and not os.path.isdir(out_dir):
            os.makedirs(out_dir, exist_ok=True)

        doc = Document()

        # Add title
        title = doc.add_heading(os.path.splitext(os.path.basename(input_pdf))[0], 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        with pdfplumber.open(input_pdf) as pdf:
            total_pages = len(pdf.pages)
            print(f"Processing PDF with {total_pages} pages...", file=sys.stderr)

            for page_num, page in enumerate(pdf.pages, start=1):
                print(f"Processing page {page_num} of {total_pages}...", file=sys.stderr)

                # Extract tables from the page
                tables = page.extract_tables()

                if tables:
                    # Process tables
                    for table_idx, table_data in enumerate(table_data for table_data in tables if table_data):
                        if not table_data or not any(table_data):
                            continue

                        # Calculate optimal column count
                        max_cols = 0
                        for row in table_data:
                            if row:
                                max_cols = max(max_cols, len([c for c in row if c is not None and str(c).strip()]))

                        if max_cols == 0:
                            continue

                        # Add table to Word document
                        word_table = doc.add_table(
                            rows=len(table_data),
                            cols=max_cols
                        )
                        word_table.style = 'Table Grid'

                        # Fill table cells
                        for r_idx, row in enumerate(table_data):
                            for c_idx, cell_text in enumerate(row):
                                if c_idx < max_cols:
                                    cell = word_table.cell(r_idx, c_idx)
                                    # Clean cell text
                                    clean_text = str(cell_text).strip() if cell_text else ''
                                    cell.text = clean_text
                                    # Add basic formatting
                                    for paragraph in cell.paragraphs:
                                        paragraph.style = 'Normal'

                        doc.add_paragraph()  # spacing after table
                        print(f"  Extracted table with {len(table_data)} rows", file=sys.stderr)
                else:
                    # Non-table pages — extract text with layout awareness
                    text = page.extract_text(layout=True)
                    if text and text.strip():
                        # Clean and add text
                        lines = text.split('\n')
                        for line in lines:
                            line = line.strip()
                            if line:
                                doc.add_paragraph(line)

                # Add page break between pages (except for the last page)
                if page_num < total_pages:
                    doc.add_page_break()

        # Save the document
        doc.save(output_docx)

        if os.path.exists(output_docx):
            print("SUCCESS")
            sys.exit(0)
        else:
            print("ERROR: Output file was not created.", file=sys.stderr)
            sys.exit(1)

    except Exception as e:
        print(f"ERROR: {str(e)}", file=sys.stderr)
        import traceback
        traceback.print_exc(file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    if len(sys.argv) != 3:
        print("ERROR: Usage: pdf_to_word.py <input_pdf> <output_docx>", file=sys.stderr)
        sys.exit(1)
    main(sys.argv[1], sys.argv[2])
